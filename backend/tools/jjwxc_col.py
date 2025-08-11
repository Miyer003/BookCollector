import os
import time
import random
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import json
from datetime import datetime
import urllib.parse

COOKIE_FILE = "my_cookie.txt"

class JJWXCBackupTool:
    def __init__(self):
        """
        åˆå§‹åŒ–å¤‡ä»½å·¥å…·
        
        åŠŸèƒ½ï¼š
        - åˆ›å»ºè¾“å‡ºç›®å½•ç»“æ„ (backup/YYYYMMDD_HHMMSS/)
        - åˆå§‹åŒ–HTTPä¼šè¯å’Œè¯·æ±‚å¤´
        - åŠ è½½Cookieæ–‡ä»¶å¹¶è§£æè®¤è¯ä¿¡æ¯
        - é…ç½®ç½‘ç»œé‡è¯•ç­–ç•¥
        """
        # åˆ›å»ºè¾“å‡ºç›®å½• - ä½¿ç”¨timestampç¡®ä¿å”¯ä¸€æ€§
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        self.output_dir = os.path.join("backup", timestamp)
        os.makedirs(self.output_dir, exist_ok=True)
        print(f"è¾“å‡ºç›®å½•: {self.output_dir}")
        
        # è®¾ç½®HTTPä¼šè¯ - ä¿æŒCookieå’Œè¿æ¥å¤ç”¨
        self.session = requests.Session()
        self.headers = self.get_default_headers()
        
        # åˆå§‹åŒ–ä½œè€…åå°URL
        self.author_backend_url = None
        
        # åŠ è½½å¹¶è§£æCookieæ–‡ä»¶
        cookie_count = self.load_cookie()
        print(f"å·²è®¾ç½® {cookie_count} ä¸ªCookieå‚æ•°")
        
        # è®¾ç½®è¯·æ±‚é‡è¯•ç­–ç•¥ - åº”å¯¹ç½‘ç»œæ³¢åŠ¨
        self.session.mount('https://', requests.adapters.HTTPAdapter(
            max_retries=3,
            pool_connections=10,
            pool_maxsize=20
        ))

    def get_default_headers(self):
        """
        è·å–é»˜è®¤HTTPè¯·æ±‚å¤´
        
        è¿”å›ï¼š
            dict: åŒ…å«User-Agentã€Acceptç­‰æ ‡å‡†æµè§ˆå™¨è¯·æ±‚å¤´çš„å­—å…¸
            
        ç”¨é€”ï¼š
            æ¨¡æ‹ŸçœŸå®æµè§ˆå™¨è®¿é—®ï¼Œé¿å…è¢«ç½‘ç«™åçˆ¬è™«ç­–ç•¥æ‹¦æˆª
        """
        return {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
            'Connection': 'keep-alive',
            'Referer': 'https://www.jjwxc.net/',
            'Cache-Control': 'no-cache',
            'Pragma': 'no-cache'
        }
    
    def decode_unicode_escape(self, s):
        """
        è§£ç Unicodeè½¬ä¹‰åºåˆ—
        
        å‚æ•°ï¼š
            s (str): åŒ…å«%uXXXXæ ¼å¼Unicodeè½¬ä¹‰çš„å­—ç¬¦ä¸²
            
        è¿”å›ï¼š
            str: è§£ç åçš„å­—ç¬¦ä¸²
            
        è¯´æ˜ï¼š
            æ™‹æ±ŸCookieä¸­å¯èƒ½åŒ…å«%uæ ¼å¼çš„Unicodeè½¬ä¹‰å­—ç¬¦
            éœ€è¦ç‰¹æ®Šå¤„ç†æ‰èƒ½æ­£ç¡®è§£æä¸­æ–‡å­—ç¬¦
        """
        def replace_unicode(match):
            return chr(int(match.group(1), 16))
        return re.sub(r'%u([0-9a-fA-F]{4})', replace_unicode, s)
    
    def load_cookie(self):
        """
        ä»æ–‡ä»¶åŠ è½½å¹¶è§£æCookie
        
        è¿”å›ï¼š
            int: æˆåŠŸè§£æçš„Cookieæ•°é‡
            
        åŠŸèƒ½ï¼š
        1. è¯»å–my_cookie.txtæ–‡ä»¶å†…å®¹
        2. å¤„ç†Unicodeè½¬ä¹‰åºåˆ—(%uXXXX)
        3. æ™ºèƒ½è§£æå¤æ‚Cookieæ ¼å¼ï¼ˆåŒ…æ‹¬JSONå€¼ï¼‰
        4. è®¾ç½®åˆ°HTTPä¼šè¯ä¸­
        
        Cookieæ ¼å¼æ”¯æŒï¼š
        - æ ‡å‡†é”®å€¼å¯¹: key=value; key2=value2
        - JSONå€¼: key={"json":"value"}  
        - URLç¼–ç å€¼: key=%E4%B8%AD%E6%96%87
        - Unicodeè½¬ä¹‰: key=%u4E2D%u6587
        
        é”™è¯¯å¤„ç†ï¼š
        - æ–‡ä»¶ä¸å­˜åœ¨ï¼šè¿”å›0
        - è§£æå¤±è´¥ï¼šè·³è¿‡è¯¥Cookieå¹¶ç»§ç»­
        - JSONéªŒè¯ï¼šç¡®ä¿JSONæ ¼å¼Cookieçš„æœ‰æ•ˆæ€§
        """
        cookie_count = 0
        if os.path.exists(COOKIE_FILE):
            try:
                with open(COOKIE_FILE, 'r', encoding='utf-8') as f:
                    # è¯»å–åŸå§‹Cookieå†…å®¹
                    raw_cookie = f.read().strip()
                    print(f"åŸå§‹Cookieå†…å®¹: {raw_cookie[:100]}...")
                    
                    # ç¬¬ä¸€æ­¥ï¼šå¤„ç†Unicodeè½¬ä¹‰åºåˆ— (%uXXXX)
                    decoded_cookie = self.decode_unicode_escape(raw_cookie)
                    
                    # ç¬¬äºŒæ­¥ï¼šæ™ºèƒ½Cookieè§£æ - å¤„ç†åŒ…å«JSONçš„å¤æ‚Cookie
                    cookies_dict = {}
                    current_pos = 0
                    
                    while current_pos < len(decoded_cookie):
                        # æ‰¾åˆ°ä¸‹ä¸€ä¸ªç­‰å·
                        eq_pos = decoded_cookie.find('=', current_pos)
                        if eq_pos == -1:
                            break
                            
                        # æå–é”®åï¼ˆå»æ‰å‰é¢çš„åˆ†å·å’Œç©ºæ ¼ï¼‰
                        key_start = current_pos
                        if decoded_cookie[current_pos:current_pos+1] == ';':
                            key_start += 1
                        key = decoded_cookie[key_start:eq_pos].strip()
                        
                        # æ‰¾å€¼çš„éƒ¨åˆ†
                        value_start = eq_pos + 1
                        
                        # å¦‚æœå€¼ä»¥ %7B æˆ– { å¼€å¤´ï¼Œå¯èƒ½æ˜¯JSONï¼Œéœ€è¦ç‰¹æ®Šå¤„ç†
                        if (decoded_cookie[value_start:value_start+3] == '%7B' or 
                            decoded_cookie[value_start:value_start+1] == '{'):
                            # JSONå€¼ï¼šæ‰¾åˆ°åŒ¹é…çš„ç»“æŸä½ç½®
                            brace_count = 0
                            value_end = value_start
                            in_string = False
                            escape_next = False
                            
                            for i in range(value_start, len(decoded_cookie)):
                                char = decoded_cookie[i]
                                
                                if escape_next:
                                    escape_next = False
                                    continue
                                    
                                if char == '\\':
                                    escape_next = True
                                    continue
                                    
                                if char == '"' and not escape_next:
                                    in_string = not in_string
                                    continue
                                    
                                if not in_string:
                                    if char == '{' or decoded_cookie[i:i+3] == '%7B':
                                        brace_count += 1
                                        if decoded_cookie[i:i+3] == '%7B':
                                            i += 2  # è·³è¿‡ %7B çš„å…¶ä½™éƒ¨åˆ†
                                    elif char == '}' or decoded_cookie[i:i+3] == '%7D':
                                        brace_count -= 1
                                        if decoded_cookie[i:i+3] == '%7D':
                                            i += 2
                                        if brace_count == 0:
                                            value_end = i + 1
                                            break
                                    elif char == ';' and brace_count == 0:
                                        value_end = i
                                        break
                                        
                                value_end = i + 1
                        else:
                            # æ™®é€šå€¼ï¼šæ‰¾åˆ°ä¸‹ä¸€ä¸ªåˆ†å·æˆ–å­—ç¬¦ä¸²ç»“å°¾
                            value_end = decoded_cookie.find(';', value_start)
                            if value_end == -1:
                                value_end = len(decoded_cookie)
                        
                        # æå–å€¼
                        value = decoded_cookie[value_start:value_end].strip()
                        
                        # å¤„ç†å€¼çš„URLè§£ç 
                        if key and value:
                            try:
                                # å¯¹äºJSONæ ¼å¼çš„Cookieå€¼ï¼Œè°¨æ…å¤„ç†URLè§£ç 
                                if value.startswith('%7B') or value.startswith('{'):
                                    # å°è¯•è§£ç ï¼Œä½†å¦‚æœå¤±è´¥å°±ä¿æŒåŸæ ·
                                    try:
                                        decoded_value = urllib.parse.unquote(value)
                                        # éªŒè¯è§£ç åçš„JSONæ˜¯å¦æœ‰æ•ˆ
                                        if decoded_value.startswith('{') and decoded_value.endswith('}'):
                                            json.loads(decoded_value)  # éªŒè¯JSONæ ¼å¼
                                    except:
                                        decoded_value = value  # è§£ç å¤±è´¥ï¼Œä¿æŒåŸæ ·
                                else:
                                    # æ™®é€šå€¼ï¼Œç›´æ¥è§£ç 
                                    decoded_value = urllib.parse.unquote(value)
                                
                                cookies_dict[key] = decoded_value
                                self.session.cookies.set(key, decoded_value)
                                print(f"  â†’ è®¾ç½®Cookie: {key}={decoded_value[:50]}...")
                                cookie_count += 1
                                
                            except Exception as decode_error:
                                print(f"  Ã— è·³è¿‡æ— æ•ˆCookie: {key} (è§£ç é”™è¯¯: {decode_error})")
                        
                        # ç§»åŠ¨åˆ°ä¸‹ä¸€ä¸ªCookie
                        current_pos = value_end
                        if current_pos < len(decoded_cookie) and decoded_cookie[current_pos] == ';':
                            current_pos += 1
                    
                    print(f"æˆåŠŸè§£æ {cookie_count} ä¸ªCookie")
                    return cookie_count
                    
            except Exception as e:
                print(f"Cookieæ–‡ä»¶è§£æé”™è¯¯: {e}")
        
        return 0
    
    def check_login(self):
        """æ£€æŸ¥ç™»å½•çŠ¶æ€"""
        self.author_backend_url = "https://my.jjwxc.net/backend/oneauthor_login.php"
        try:
            print("æ­£åœ¨æ£€æŸ¥ç™»å½•çŠ¶æ€...")
            response = self.session.get(self.author_backend_url, headers=self.headers, timeout=15)
            response.encoding = 'gb18030'
            html = response.text

            # åˆ¤æ–­é¡µé¢æ˜¯å¦åŒ…å«ç™»å½•æç¤º
            if "æ™‹æ±Ÿæ–‡å­¦åŸ" in html:
                print("ç™»å…¥æˆåŠŸ")
                return True
            elif "è¯·ç™»å½•" in html or "ç™»å½•æ™‹æ±Ÿä½œè€…åå°" in html or "è´¦å·" in html:
                print("æœªç™»å½•æ™‹æ±Ÿä½œè€…åå°ï¼Œè¯·æ£€æŸ¥Cookieæ˜¯å¦æœ‰æ•ˆ")
                return False
        except Exception as e:
            print(f"æ£€æŸ¥ç™»å½•çŠ¶æ€æ—¶å‡ºé”™: {e}")
            return False
        
    def get_novel_list(self):
        """è·å–ä½œè€…ä½œå“åˆ—è¡¨"""
        author_url = "https://my.jjwxc.net/backend/oneauthor_login.php"
        
        try:
            print(f"è·å–ä½œå“åˆ—è¡¨: {author_url}")
            response = self.session.get(author_url, headers=self.headers, timeout=20)
            
            # è®¾ç½®æ­£ç¡®çš„ç¼–ç 
            response.encoding = 'gb18030'  # æ™‹æ±Ÿä½¿ç”¨gb18030ç¼–ç 
            
            soup = BeautifulSoup(response.content, 'html.parser', from_encoding='gb18030')
            novels = []
            
            # ä¿å­˜é¡µé¢ç”¨äºè°ƒè¯•ï¼ˆå·²ç¦ç”¨ï¼Œå¦‚éœ€è°ƒè¯•è¯·å–æ¶ˆæ³¨é‡Šï¼‰
            # with open(os.path.join(self.output_dir, "novel_list.html"), "w", encoding="utf-8") as f:
            #     f.write(response.text)
            # print("ä½œå“åˆ—è¡¨é¡µé¢å·²ä¿å­˜: novel_list.html")
            
            # æŸ¥æ‰¾ä½œå“ç®¡ç†é“¾æ¥
            # åœ¨æ™‹æ±Ÿåå°ï¼Œä½œå“ç®¡ç†é“¾æ¥çš„æ ¼å¼æ˜¯: managenovel.php?novelid=XXXXX
            novel_links = soup.find_all('a', href=lambda x: x and 'managenovel.php?novelid=' in x)
            
            if novel_links:
                print(f"æ‰¾åˆ° {len(novel_links)} ä¸ªä½œå“ç®¡ç†é“¾æ¥")
                
                for link in novel_links:
                    # æå–ä½œå“ID
                    href = link['href']
                    novel_id_match = re.search(r'novelid=(\d+)', href)
                    if novel_id_match:
                        novel_id = novel_id_match.group(1)
                        
                        # æŸ¥æ‰¾ä½œå“æ ‡é¢˜ï¼ˆåœ¨åŒä¸€è¡Œçš„å…¶ä»–ä½ç½®ï¼‰
                        row = link.find_parent('tr')
                        if row:
                            # æŸ¥æ‰¾ä½œå“æ ‡é¢˜é“¾æ¥ï¼ˆæŒ‡å‘onebook.phpçš„é“¾æ¥ï¼‰
                            title_link = row.find('a', href=lambda x: x and f'onebook.php?novelid={novel_id}' in x)
                            if title_link:
                                title = title_link.get_text(strip=True)
                                
                        # æå–å…¶ä»–ä¿¡æ¯
                        cells = row.find_all('td')
                        if len(cells) >= 10:
                            try:
                                # æ ¹æ®HTMLç»“æ„æå–ä¿¡æ¯
                                category = cells[2].get_text(strip=True) if len(cells) > 2 else "æœªçŸ¥"
                                subcategory = cells[3].get_text(strip=True) if len(cells) > 3 else "æœªçŸ¥"
                                chapter_count = cells[5].get_text(strip=True) if len(cells) > 5 else "0"
                                word_count = cells[6].get_text(strip=True) if len(cells) > 6 else "0"
                                status = cells[12].get_text(strip=True) if len(cells) > 12 else "æœªçŸ¥"
                                
                                novels.append({
                                    'id': novel_id,
                                    'title': title,
                                    'link': href,
                                    'view_link': title_link['href'],
                                    'status': status,
                                    'word_count': word_count,
                                    'chapter_count': chapter_count,
                                    'category': f"{category}-{subcategory}"
                                })
                                
                            except Exception as e:
                                print(f"è§£æä½œå“ä¿¡æ¯å‡ºé”™ {novel_id}: {e}")
                                novels.append({
                                    'id': novel_id,
                                    'title': title,
                                    'link': href,
                                    'view_link': title_link['href'],
                                    'status': "æœªçŸ¥",
                                    'word_count': "æœªçŸ¥",
                                    'chapter_count': "æœªçŸ¥",
                                    'category': "æœªçŸ¥"
                                })
                
                print(f"æˆåŠŸè§£æ {len(novels)} éƒ¨ä½œå“")
                return novels
            else:
                print("æœªæ‰¾åˆ°ä½œå“ç®¡ç†é“¾æ¥")
                # å¤‡ç”¨æ–¹æ³•ï¼šæŸ¥æ‰¾onebook.phpé“¾æ¥
                onebook_links = soup.find_all('a', href=lambda x: x and 'onebook.php?novelid=' in x)
                if onebook_links:
                    print(f"æ‰¾åˆ° {len(onebook_links)} ä¸ªä½œå“é˜…è¯»é“¾æ¥")
                    for idx, link in enumerate(onebook_links):
                        href = link['href']
                        novel_id_match = re.search(r'novelid=(\d+)', href)
                        if novel_id_match:
                            novel_id = novel_id_match.group(1)
                            title = link.get_text(strip=True)
                            
                            novels.append({
                                'id': novel_id,
                                'title': title,
                                'link': f"//my.jjwxc.net/backend/managenovel.php?novelid={novel_id}",
                                'view_link': href,
                                'status': "æœªçŸ¥",
                                'word_count': "æœªçŸ¥",
                                'chapter_count': "æœªçŸ¥",
                                'category': "æœªçŸ¥"
                            })
                    
                    return novels
                else:
                    print("ä¹Ÿæœªæ‰¾åˆ°ä½œå“é˜…è¯»é“¾æ¥")
                    return []
            
        except Exception as e:
            print(f"è·å–ä½œå“åˆ—è¡¨å‡ºé”™: {str(e)}")
            return []
    
    def get_intro_from_backend(self, novel_id):
        """ä»ä½œè€…åå°è·å–ä½œå“ç®€ä»‹"""
        try:
            backend_url = f"https://my.jjwxc.net/backend/managenovel.php?novelid={novel_id}"
            print(f"è®¿é—®åå°ç« èŠ‚ç®¡ç†é¡µé¢: {backend_url}")
            headers = self.headers.copy()
            headers['Referer'] = 'https://my.jjwxc.net/backend/'
            response = self.session.get(backend_url, headers=headers, timeout=30)
            response.encoding = 'gb18030'
            
            soup = BeautifulSoup(response.content, 'html.parser', from_encoding='gb18030')
            novel_intro = ""
            intro_textarea = soup.find('textarea', {'id': 'novelintro'})
            if intro_textarea:
                novel_intro = intro_textarea.get_text(strip=True)
                print(f"è·å–åˆ°ä½œå“ç®€ä»‹: {len(novel_intro)} å­—ç¬¦")
            return novel_intro
        except Exception as e:
            print(f"è·å–ä½œå“ç®€ä»‹å¤±è´¥: {e}")
            return ""

    def get_chapters(self, novel_link):
        """
        è·å–ä½œå“çš„å®Œæ•´ç« èŠ‚åˆ—è¡¨ï¼ˆç»Ÿä¸€åå°æ–¹æ¡ˆï¼‰
        
        å‚æ•°ï¼š
            novel_link (str): ä½œå“ç®¡ç†é¡µé¢é“¾æ¥
            
        è¿”å›ï¼š
            list: ç« èŠ‚ä¿¡æ¯åˆ—è¡¨ï¼Œæ¯ä¸ªå…ƒç´ åŒ…å«ï¼š
                - id: ç« èŠ‚ID
                - title: ç« èŠ‚æ ‡é¢˜  
                - link: åå°ç¼–è¾‘é¡µé¢é“¾æ¥ï¼ˆç»Ÿä¸€æ ¼å¼ï¼‰
                - chapter_number: ç« èŠ‚ç¼–å·
                - is_vip: æ˜¯å¦VIPç« èŠ‚
                
        æ–°æ–¹æ¡ˆè¯´æ˜ï¼š
        - ç»Ÿä¸€ä½¿ç”¨åå°ç¼–è¾‘é¡µé¢è·å–æ‰€æœ‰ç« èŠ‚å†…å®¹
        - é¿å…å‰å°é¡µé¢ç»“æ„å˜åŒ–å¯¼è‡´çš„è§£æé—®é¢˜
        - å…è´¹å’ŒVIPç« èŠ‚ä½¿ç”¨ç›¸åŒçš„è·å–é€»è¾‘
        - æé«˜ä»£ç çš„ç¨³å®šæ€§å’Œå¯ç»´æŠ¤æ€§
        """
        if not novel_link:
            return []
            
        try:
            # æå–novelid
            novel_id_match = re.search(r'novelid=(\d+)', novel_link)
            if not novel_id_match:
                print(f"æ— æ³•ä»é“¾æ¥ä¸­æå–ä½œå“ID: {novel_link}")
                return []
            novel_id = novel_id_match.group(1)
            backend_url = f"https://my.jjwxc.net/backend/managenovel.php?novelid={novel_id}"
            print(f"è·å–æ‰€æœ‰ç« èŠ‚åˆ—è¡¨: {backend_url}")
            response = self.session.get(backend_url, headers=self.headers, timeout=30)
            response.encoding = 'gb18030'
            
            soup = BeautifulSoup(response.content, 'html.parser', from_encoding='gb18030')
            chapters = []
            
            # æ–°çš„ç« èŠ‚è§£æç­–ç•¥ï¼šé€šè¿‡å¤šç§æ–¹æ³•ç»„åˆæŸ¥æ‰¾
            # æ–¹æ³•1ï¼šæŸ¥æ‰¾æ‰€æœ‰åŒ…å«chapteridçš„inputï¼Œé€šè¿‡å…¶çˆ¶å…ƒç´ å®šä½ç« èŠ‚è¡Œ
            chapter_inputs = soup.find_all('input', {'name': 'chapterid'})
            print(f"æ‰¾åˆ° {len(chapter_inputs)} ä¸ªç« èŠ‚IDè¾“å…¥æ¡†")
            
            # è¿‡æ»¤æ‰è¡¨å•æäº¤ç”¨çš„hidden inputï¼ˆé€šå¸¸åœ¨formä¸­ä¸”valueè¾ƒå¤§ï¼‰
            valid_chapter_inputs = []
            for input_elem in chapter_inputs:
                # æ£€æŸ¥æ˜¯å¦åœ¨è¡¨å•ä¸­ä¸”valueå¾ˆå¤§ï¼ˆå¯èƒ½æ˜¯ä¸‹ä¸€ç« èŠ‚çš„IDï¼‰
                form_parent = input_elem.find_parent('form')
                if form_parent and input_elem.get('type') == 'hidden':
                    # è¿™å¯èƒ½æ˜¯è¡¨å•æäº¤ç”¨çš„inputï¼Œè·³è¿‡
                    continue
                valid_chapter_inputs.append(input_elem)
            
            print(f"æœ‰æ•ˆç« èŠ‚è¾“å…¥æ¡†: {len(valid_chapter_inputs)}")
            
            # å¦‚æœæ²¡æœ‰æ‰¾åˆ°æœ‰æ•ˆçš„ç« èŠ‚è¾“å…¥æ¡†ï¼Œå°è¯•å…¶ä»–æ–¹æ³•
            if not valid_chapter_inputs:
                print("å°è¯•å¤‡ç”¨æ–¹æ³•ï¼šæŸ¥æ‰¾ç« èŠ‚é“¾æ¥")
                # æ–¹æ³•2ï¼šæŸ¥æ‰¾æ‰€æœ‰æŒ‡å‘ç« èŠ‚çš„é“¾æ¥
                all_links = soup.find_all('a', href=True)
                chapter_links = []
                for link in all_links:
                    href = link.get('href', '')
                    if 'onebook' in href and ('novelid=' in href or 'chapterid=' in href):
                        chapter_links.append(link)
                
                print(f"æ‰¾åˆ° {len(chapter_links)} ä¸ªç« èŠ‚é“¾æ¥")
                
                # ä»é“¾æ¥ä¸­æå–ç« èŠ‚ä¿¡æ¯
                for link in chapter_links:
                    href = link.get('href')
                    chapter_id_match = re.search(r'chapterid=(\d+)', href)
                    
                    if not chapter_id_match:
                        continue
                    
                    chapter_id = chapter_id_match.group(1)
                    title = link.get_text(strip=True)
                    
                    # åˆ¤æ–­æ˜¯å¦VIPç« èŠ‚
                    is_vip = 'onebook_vip.php' in href or '[VIP]' in title
                    
                    # å°è¯•ä»é“¾æ¥çš„çˆ¶å…ƒç´ ä¸­è·å–ç« èŠ‚ç¼–å·
                    parent_tr = link.find_parent('tr')
                    chapter_number = len(chapters) + 1  # é»˜è®¤æŒ‰é¡ºåºç¼–å·
                    
                    if parent_tr:
                        tds = parent_tr.find_all('td')
                        for td in tds:
                            text = td.get_text(strip=True)
                            # æŸ¥æ‰¾æ•°å­—ç¼–å·
                            number_match = re.search(r'^(\d+)$', text)
                            if number_match:
                                chapter_number = int(number_match.group(1))
                                break
                    
                    # æ„å»ºç»Ÿä¸€çš„åå°ç¼–è¾‘é“¾æ¥
                    edit_link = f"https://my.jjwxc.net/backend/chaptermodify.php?novelid={novel_id}&chapterid={chapter_id}"
                    
                    chapters.append({
                        'id': chapter_id,
                        'title': title,
                        'link': edit_link,  # ç»Ÿä¸€ä½¿ç”¨åå°ç¼–è¾‘é“¾æ¥
                        'chapter_number': chapter_number,
                        'is_vip': is_vip
                    })
            
            else:
                # ä½¿ç”¨æœ‰æ•ˆçš„ç« èŠ‚è¾“å…¥æ¡†è§£æç« èŠ‚
                for input_elem in valid_chapter_inputs:
                    chapter_id = input_elem.get('value')
                    if not chapter_id:
                        continue
                    
                    # æŸ¥æ‰¾å…³è”çš„ç« èŠ‚è¡Œ
                    parent_tr = input_elem.find_parent('tr')
                    if not parent_tr:
                        continue
                    
                    # æŸ¥æ‰¾ç« èŠ‚æ ‡é¢˜é“¾æ¥
                    title_link = parent_tr.find('a', href=True)
                    if not title_link:
                        continue
                        
                    title = title_link.get_text(strip=True)
                    href = title_link.get('href')
                    
                    # åˆ¤æ–­æ˜¯å¦VIPç« èŠ‚
                    is_vip = 'onebook_vip.php' in href or '[VIP]' in title
                    
                    # æå–ç« èŠ‚ç¼–å·
                    chapter_number = len(chapters) + 1  # é»˜è®¤ç¼–å·
                    tds = parent_tr.find_all('td')
                    if len(tds) > 1:
                        try:
                            chapter_num_text = tds[1].get_text(strip=True)
                            chapter_number = int(chapter_num_text)
                        except (ValueError, IndexError):
                            pass
                    
                    # æ„å»ºç»Ÿä¸€çš„åå°ç¼–è¾‘é“¾æ¥
                    edit_link = f"https://my.jjwxc.net/backend/chaptermodify.php?novelid={novel_id}&chapterid={chapter_id}"
                    
                    chapters.append({
                        'id': chapter_id,
                        'title': title,
                        'link': edit_link,  # ç»Ÿä¸€ä½¿ç”¨åå°ç¼–è¾‘é“¾æ¥
                        'chapter_number': chapter_number,
                        'is_vip': is_vip
                    })
            
            # å¦‚æœå¸¸è§„æ–¹æ³•éƒ½å¤±è´¥ï¼Œå°è¯•é€šè¿‡æœ€å¤§ç« èŠ‚å·ç”Ÿæˆç« èŠ‚åˆ—è¡¨
            if not chapters:
                print("å°è¯•æœ€ç»ˆæ–¹æ¡ˆï¼šé€šè¿‡æœ€å¤§ç« èŠ‚å·ç”Ÿæˆç« èŠ‚åˆ—è¡¨")
                
                # æŸ¥æ‰¾æœ€å¤§ç« èŠ‚å·æç¤º
                max_chapter_hints = [
                    soup.find(text=re.compile(r'å·²æ›´æ–°è‡³ç¬¬(\d+)ç« ')),
                    soup.find(text=re.compile(r'ç¬¬(\d+)ç« ', re.I))
                ]
                
                max_chapter_num = 0
                for hint in max_chapter_hints:
                    if hint:
                        match = re.search(r'ç¬¬(\d+)ç« ', str(hint))
                        if match:
                            chapter_num = int(match.group(1))
                            max_chapter_num = max(max_chapter_num, chapter_num)
                
                # ä¹Ÿå¯ä»¥æŸ¥çœ‹placeholderä¸­çš„ç« èŠ‚å·ï¼ˆä¸‹ä¸€ç« èŠ‚å·ï¼‰
                placeholders = soup.find_all('input', {'placeholder': re.compile(r'ç¬¬(\d+)ç« ')})
                for placeholder in placeholders:
                    placeholder_text = placeholder.get('placeholder', '')
                    match = re.search(r'ç¬¬(\d+)ç« ', placeholder_text)
                    if match:
                        next_chapter_num = int(match.group(1))
                        max_chapter_num = max(max_chapter_num, next_chapter_num - 1)
                
                print(f"æ£€æµ‹åˆ°æœ€å¤§ç« èŠ‚å·: {max_chapter_num}")
                
                if max_chapter_num > 0:
                    print(f"ç”Ÿæˆ 1-{max_chapter_num} ç« èŠ‚åˆ—è¡¨")
                    for chapter_num in range(1, max_chapter_num + 1):
                        # æ„å»ºç»Ÿä¸€çš„åå°ç¼–è¾‘é“¾æ¥
                        edit_link = f"https://my.jjwxc.net/backend/chaptermodify.php?novelid={novel_id}&chapterid={chapter_num}"
                        
                        chapters.append({
                            'id': str(chapter_num),
                            'title': f"ç¬¬{chapter_num}ç« ",  # ä¸´æ—¶æ ‡é¢˜ï¼Œåç»­å¯ä»¥ä»ç¼–è¾‘é¡µé¢è·å–
                            'link': edit_link,
                            'chapter_number': chapter_num,
                            'is_vip': False  # æš‚æ—¶æ ‡è®°ä¸ºå…è´¹ï¼Œå®é™…ç±»å‹ä¼šåœ¨è·å–å†…å®¹æ—¶ç¡®å®š
                        })
            
            # æŒ‰ç« èŠ‚ç¼–å·æ’åº
            chapters.sort(key=lambda x: x['chapter_number'])
            
            vip_count = sum(1 for c in chapters if c['is_vip'])
            free_count = len(chapters) - vip_count
            print(f"æˆåŠŸè§£æ {len(chapters)} ä¸ªç« èŠ‚ï¼Œå…¶ä¸­å…è´¹ç« èŠ‚æ•°é‡ï¼š{free_count}ï¼ŒVIPç« èŠ‚æ•°é‡ï¼š{vip_count}")
            return chapters
        except Exception as e:
            print(f"è·å–ç« èŠ‚åˆ—è¡¨å‡ºé”™: {str(e)}")
            return []
    
    def get_chapter_content(self, chapter_link, is_vip=False):
        """
        è·å–ç« èŠ‚å†…å®¹ï¼ˆç»Ÿä¸€åå°æ–¹æ¡ˆï¼‰
        
        å‚æ•°ï¼š
            chapter_link (str): åå°ç¼–è¾‘é¡µé¢é“¾æ¥
            is_vip (bool): æ˜¯å¦ä¸ºVIPç« èŠ‚ï¼ˆä¿ç•™å‚æ•°ï¼Œä½†ä¸å†å½±å“å¤„ç†é€»è¾‘ï¼‰
            
        è¿”å›ï¼š
            str: ç« èŠ‚å®Œæ•´å†…å®¹ï¼ˆåŒ…å«æ­£æ–‡å’Œä½œè€…æœ‰è¯è¯´ï¼‰
            
        ç»Ÿä¸€å¤„ç†æ–¹æ¡ˆï¼š
        - æ‰€æœ‰ç« èŠ‚éƒ½é€šè¿‡åå°ç¼–è¾‘é¡µé¢è·å–å†…å®¹
        - ä»textareaå…ƒç´ è·å–åŸå§‹æœªåŠ å¯†å†…å®¹
        - ç»Ÿä¸€çš„æ ¼å¼ä¿ç•™å’Œé”™è¯¯å¤„ç†é€»è¾‘
        - é¿å…å‰å°é¡µé¢ç»“æ„å˜åŒ–çš„å½±å“
        
        å¤„ç†æµç¨‹ï¼š
        1. è®¿é—®åå°ç¼–è¾‘é¡µé¢ï¼ˆchaptermodify.phpï¼‰
        2. è§£ætextareaè·å–åŸå§‹å†…å®¹ï¼š
           - name='content': ç« èŠ‚æ­£æ–‡
           - name='note': ä½œè€…æœ‰è¯è¯´
        3. æ¸…ç†HTMLå®ä½“ç¼–ç ä½†ä¿æŒæ–‡æœ¬ç»“æ„
        4. ç»„åˆæ­£æ–‡å’Œä½œè€…æœ‰è¯è¯´å†…å®¹
        
        ä¼˜åŠ¿ï¼š
        - ä»£ç é€»è¾‘ç»Ÿä¸€ï¼Œç»´æŠ¤ç®€å•
        - é¿å…å‰å°é¡µé¢è§£æé—®é¢˜
        - æ ¼å¼ä¿ç•™æ•ˆæœæ›´å¥½
        - å…è´¹å’ŒVIPç« èŠ‚ä½¿ç”¨ç›¸åŒé€»è¾‘
        """
        if not chapter_link:
            return "ç« èŠ‚é“¾æ¥æ— æ•ˆ"
            
        try:
            print(f"  è·å–ç« èŠ‚å†…å®¹ï¼ˆç»Ÿä¸€åå°æ–¹æ¡ˆï¼‰...")
            
            # å¦‚æœä¼ å…¥çš„ä¸æ˜¯åå°ç¼–è¾‘é“¾æ¥ï¼Œéœ€è¦è½¬æ¢
            if 'chaptermodify.php' not in chapter_link:
                # ä»åŸå§‹é“¾æ¥æå–novelidå’Œchapterid
                novel_id_match = re.search(r'novelid=(\d+)', chapter_link)
                chapter_id_match = re.search(r'chapterid=(\d+)', chapter_link)
                
                if not novel_id_match or not chapter_id_match:
                    return "æ— æ³•ä»é“¾æ¥ä¸­æå–ç« èŠ‚ä¿¡æ¯"
                
                novel_id = novel_id_match.group(1)
                chapter_id = chapter_id_match.group(1)
                
                # æ„å»ºåå°ç¼–è¾‘é¡µé¢é“¾æ¥
                edit_url = f"https://my.jjwxc.net/backend/chaptermodify.php?novelid={novel_id}&chapterid={chapter_id}"
            else:
                edit_url = chapter_link
            
            # è®¾ç½®è¯·æ±‚å¤´
            headers = self.headers.copy()
            headers['Referer'] = f'https://my.jjwxc.net/backend/managenovel.php'
            
            # è®¿é—®åå°ç¼–è¾‘é¡µé¢
            response = self.session.get(edit_url, headers=headers, timeout=30)
            response.encoding = 'gb18030'
            soup = BeautifulSoup(response.content, 'html.parser', from_encoding='gb18030')
            
            main_content = ""
            author_notes = ""
            
            # ä»ç¼–è¾‘é¡µé¢çš„textareaè·å–æ­£æ–‡å†…å®¹
            chapterbody_textarea = soup.find('textarea', {'name': 'content'})
            if chapterbody_textarea:
                # è·å–åŸå§‹æ–‡æœ¬å†…å®¹ï¼Œä¿ç•™æ‰€æœ‰æ ¼å¼
                main_content = chapterbody_textarea.string or chapterbody_textarea.get_text()
                # å¦‚æœæ²¡æœ‰å†…å®¹ï¼Œå°è¯•ä»textareaå†…éƒ¨è·å–
                if not main_content.strip():
                    main_content = ''.join(str(content) for content in chapterbody_textarea.contents)
                
                # åªæ¸…ç†HTMLå®ä½“ç¼–ç ï¼Œä¿ç•™æ‰€æœ‰æ¢è¡Œå’Œç©ºè¡Œ
                main_content = main_content.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
                main_content = main_content.replace('&quot;', '"').replace('&#039;', "'")
                main_content = main_content.replace('&nbsp;', ' ')  # å¤„ç†éæ–­è¡Œç©ºæ ¼
            
            # è·å–ä½œè€…æœ‰è¯è¯´
            authornote_textarea = soup.find('textarea', {'name': 'note'})
            if authornote_textarea:
                # è·å–åŸå§‹æ–‡æœ¬å†…å®¹ï¼Œä¿ç•™æ‰€æœ‰æ ¼å¼
                author_notes = authornote_textarea.string or authornote_textarea.get_text()
                # å¦‚æœæ²¡æœ‰å†…å®¹ï¼Œå°è¯•ä»textareaå†…éƒ¨è·å–
                if not author_notes.strip():
                    author_notes = ''.join(str(content) for content in authornote_textarea.contents)
                
                # åªæ¸…ç†HTMLå®ä½“ç¼–ç ï¼Œä¿ç•™æ‰€æœ‰æ¢è¡Œå’Œç©ºè¡Œ
                author_notes = author_notes.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
                author_notes = author_notes.replace('&quot;', '"').replace('&#039;', "'")
                author_notes = author_notes.replace('&nbsp;', ' ')  # å¤„ç†éæ–­è¡Œç©ºæ ¼
            
            # ç»„åˆå†…å®¹
            result_parts = []
            if main_content and len(main_content.strip()) > 20:
                result_parts.append(main_content)
            if author_notes and len(author_notes.strip()) > 10:
                result_parts.append('\n\nã€ä½œè€…æœ‰è¯è¯´ã€‘\n')
                result_parts.append(author_notes)
            
            if result_parts:
                result = ''.join(result_parts)
                if len(result.strip()) > 30:
                    return result
            
            return "å†…å®¹è·å–å¤±è´¥ï¼šæœªæ‰¾åˆ°æœ‰æ•ˆå†…å®¹"
            
        except Exception as e:
            print(f"  ç« èŠ‚å†…å®¹è·å–å‡ºé”™: {str(e)}")
            return f"å†…å®¹è·å–å¤±è´¥ï¼š{str(e)}"

    def create_docx_with_realtime_save(self, novel, chapters):
        """
        åˆ›å»ºDOCXæ–‡æ¡£å¹¶å®æ—¶ä¿å­˜ç« èŠ‚å†…å®¹
        
        å‚æ•°ï¼š
            novel (dict): ä½œå“ä¿¡æ¯å­—å…¸
            chapters (list): ç« èŠ‚åˆ—è¡¨
            
        åŠŸèƒ½ç‰¹æ€§ï¼š
        1. æ–‡æ¡£ç»“æ„åˆ›å»ºï¼š
           - ä½œå“æ ‡é¢˜ï¼ˆ0çº§æ ‡é¢˜ï¼Œå±…ä¸­ï¼‰
           - ä½œå“ä¿¡æ¯ï¼ˆIDã€å­—æ•°ã€çŠ¶æ€ï¼‰
           - ä½œå“ç®€ä»‹ï¼ˆä»åå°è·å–ï¼‰
           - åˆ†é¡µç¬¦åˆ†éš”
           
        2. ç« èŠ‚å¤„ç†ï¼š
           - é€ç« èŠ‚è·å–å’Œæ·»åŠ å†…å®¹
           - å®æ—¶ä¿å­˜ï¼ˆæ¯ç« èŠ‚ä¿å­˜ä¸€æ¬¡ï¼‰
           - ç« èŠ‚æ ‡é¢˜æ ¼å¼åŒ–ï¼ˆç¬¬Xç«  æ ‡é¢˜ï¼‰
           - ç« èŠ‚é—´åˆ†éš”ç¬¦
           
        3. å†…å®¹æ ¼å¼åŒ–ï¼š
           - è°ƒç”¨_add_content_to_docå¤„ç†æ­£æ–‡
           - è‡ªåŠ¨åˆ†ç¦»ä½œè€…æœ‰è¯è¯´
           - ä¿ç•™åŸå§‹æ¢è¡Œå’Œç©ºè¡Œ
           - é”™è¯¯ç« èŠ‚æ ‡çº¢æ˜¾ç¤º
           
        4. å®æ—¶ä¿å­˜æœºåˆ¶ï¼š
           - åˆ›å»ºåˆå§‹æ–‡æ¡£ç»“æ„ç«‹å³ä¿å­˜
           - æ¯æ·»åŠ ä¸€ç« èŠ‚å†…å®¹åä¿å­˜
           - ç”¨æˆ·å¯éšæ—¶æ‰“å¼€æŸ¥çœ‹è¿›åº¦
           - é¿å…ç¨‹åºä¸­æ–­å¯¼è‡´æ•°æ®ä¸¢å¤±
           
        5. æ–‡ä»¶å‘½åï¼š
           - æ¸…ç†æ ‡é¢˜ä¸­çš„éæ³•å­—ç¬¦
           - ç”Ÿæˆsafeçš„æ–‡ä»¶å
           - ä¿å­˜åˆ°backup/timestamp/ç›®å½•
           
        6. é”™è¯¯å¤„ç†ï¼š
           - ç« èŠ‚è·å–å¤±è´¥ï¼šæ ‡çº¢æ˜¾ç¤ºé”™è¯¯ä¿¡æ¯
           - ç½‘ç»œå¼‚å¸¸ï¼šç»§ç»­å¤„ç†ä¸‹ä¸€ç« èŠ‚
           - æ–‡æ¡£ä¿å­˜å¼‚å¸¸ï¼šè®°å½•é”™è¯¯å¹¶ç»§ç»­
           
        7. è¿›åº¦æ˜¾ç¤ºï¼š
           - æ˜¾ç¤ºå½“å‰ç« èŠ‚è¿›åº¦ [X/æ€»æ•°]
           - é¢„ä¼°å‰©ä½™æ—¶é—´
           - ç« èŠ‚è·å–çŠ¶æ€åé¦ˆ
        """
        if not chapters:
            print(f"æ²¡æœ‰æ‰¾åˆ°ç« èŠ‚å†…å®¹ï¼Œè·³è¿‡ {novel['title']}")
            return
        
        try:
            # åˆ›å»ºWordæ–‡æ¡£
            doc = Document()
            
            # æ·»åŠ ä½œå“æ ‡é¢˜ï¼ˆæœ€é«˜çº§æ ‡é¢˜ï¼‰
            title_paragraph = doc.add_heading(novel['title'], level=0)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # æ·»åŠ ä½œå“åŸºæœ¬ä¿¡æ¯
            info_paragraph = doc.add_paragraph()
            info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            info_run = info_paragraph.add_run(
                f"ä½œå“ID: {novel['id']} | "
                f"å­—æ•°: {novel.get('word_count', 'æœªçŸ¥')} | "
                f"çŠ¶æ€: {novel.get('status', 'æœªçŸ¥')}"
            )
            info_run.font.size = Pt(10)
            
            # è·å–ä½œå“ç®€ä»‹å¹¶æ’å…¥åˆ°çŠ¶æ€ä¸‹æ–¹
            novel_intro = self.get_intro_from_backend(novel['id'])
            if novel_intro:
                intro_paragraph = doc.add_paragraph(novel_intro)
                intro_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                intro_paragraph.runs[0].font.size = Pt(11)
            
            # æ·»åŠ åˆ†é¡µç¬¦
            doc.add_page_break()
            
            # å‡†å¤‡æ–‡ä»¶åå’Œè·¯å¾„
            filename = self._clean_filename(novel['title'])
            filepath = os.path.join(self.output_dir, f"{filename}.docx")
            
            total_chapters = len(chapters)
            print(f"å¼€å§‹å¤„ç†: {novel['title']} ({total_chapters}ç« )")
            print(f"æ–‡æ¡£å°†ä¿å­˜ä¸º: {filepath}")
            
            # å…ˆä¿å­˜åˆå§‹æ–‡æ¡£ç»“æ„
            doc.save(filepath)
            print(f"âœ“ å·²åˆ›å»ºåˆå§‹æ–‡æ¡£ï¼Œå¯ä»¥æ‰“å¼€æŸ¥çœ‹")
            
            # é€ç« èŠ‚å¤„ç†å¹¶å®æ—¶ä¿å­˜
            for idx, chapter in enumerate(chapters):
                try:
                    # æ·»åŠ ç« èŠ‚æ ‡é¢˜ï¼ˆå¸¦ç« èŠ‚ç¼–å·ï¼‰
                    chapter_title = f"ç¬¬{chapter.get('chapter_number', idx+1)}ç«  {chapter['title']}"
                    doc.add_heading(chapter_title, level=1)
                    
                    # è·å–ç« èŠ‚å†…å®¹ï¼ˆç»Ÿä¸€åå°æ–¹æ¡ˆï¼‰
                    print(f"æ­£åœ¨è·å–: {chapter_title} [{idx+1}/{total_chapters}]")
                    content = self.get_chapter_content(chapter['link'])
                    
                    # æ£€æŸ¥å†…å®¹æ˜¯å¦æœ‰æ•ˆ
                    if content and not content.startswith("å†…å®¹è·å–å¤±è´¥") and not content.startswith("ç« èŠ‚é“¾æ¥æ— æ•ˆ"):
                        self._add_content_to_doc(doc, content)
                    else:
                        # å†…å®¹è·å–å¤±è´¥çš„æƒ…å†µ
                        error_paragraph = doc.add_paragraph(f"[ç« èŠ‚å†…å®¹è·å–å¤±è´¥: {content}]")
                        error_paragraph.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    
                    # æ·»åŠ ç« èŠ‚åˆ†éš”ç¬¦
                    if idx < total_chapters - 1:
                        doc.add_paragraph()
                        separator = doc.add_paragraph("â”€" * 50)
                        separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        doc.add_paragraph()
                    
                    # å®æ—¶ä¿å­˜æ–‡æ¡£
                    doc.save(filepath)
                    print(f"âœ“ å·²ä¿å­˜ [{idx+1}/{total_chapters}]")
                    
                except Exception as e:
                    print(f"å¤„ç†ç« èŠ‚å‡ºé”™: {str(e)}")
                    error_paragraph = doc.add_paragraph(f"[ç« èŠ‚å¤„ç†é”™è¯¯: {chapter['title']} - {str(e)}]")
                    error_paragraph.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    doc.save(filepath)
                
                # å»¶è¿Ÿé¿å…è¯·æ±‚è¿‡å¿«
                time.sleep(random.uniform(1.0, 2.0))
            
            print(f"âœ“ å®Œæˆä¿å­˜: {novel['title']}")
            
        except Exception as e:
            print(f"åˆ›å»ºæ–‡æ¡£å‡ºé”™: {str(e)}")
    
    def _clean_filename(self, filename):
        """æ¸…ç†æ–‡ä»¶åä¸­çš„éæ³•å­—ç¬¦"""
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        if not filename.strip() or filename.strip() == '_':
            filename = f"novel_{int(time.time())}"
        
        return filename
    
    def _add_content_to_doc(self, doc, content):
        """
        å°†ç« èŠ‚å†…å®¹æ·»åŠ åˆ°DOCXæ–‡æ¡£ä¸­ï¼ˆæ ¼å¼ä¿ç•™ç‰ˆï¼‰
        
        å‚æ•°ï¼š
            doc: python-docx Documentå¯¹è±¡
            content (str): ç« èŠ‚å®Œæ•´å†…å®¹
            
        åŠŸèƒ½ï¼š
        1. å†…å®¹åˆ†ç¦»ï¼š
           - ä»¥ã€ä½œè€…æœ‰è¯è¯´ã€‘ä¸ºåˆ†ç•Œç‚¹
           - åˆ†ç¦»æ­£æ–‡å’Œä½œè€…æœ‰è¯è¯´ä¸¤éƒ¨åˆ†
           - å»é™¤é¦–å°¾ç©ºç™½ä½†ä¿ç•™å†…éƒ¨æ ¼å¼
           
        2. æ­£æ–‡å¤„ç†ï¼š
           - æŒ‰æ¢è¡Œç¬¦(\n)åˆ†å‰²ä¸ºè¡Œ
           - æ¯è¡Œåˆ›å»ºç‹¬ç«‹æ®µè½
           - å®Œæ•´ä¿ç•™ç©ºè¡Œï¼ˆç©ºæ®µè½ï¼‰
           - ä¸åšä»»ä½•æ–‡æœ¬æ¸…ç†æˆ–å»ç©ºæ ¼
           
        3. ä½œè€…æœ‰è¯è¯´å¤„ç†ï¼š
           - æ·»åŠ è“è‰²äºŒçº§æ ‡é¢˜"ä½œè€…æœ‰è¯è¯´"
           - åŒæ ·æŒ‰è¡Œå¤„ç†å†…å®¹
           - ä¿æŒä¸æ­£æ–‡ç›¸åŒçš„æ ¼å¼å¤„ç†æ–¹å¼
           
        æ ¼å¼ä¿ç•™ç­–ç•¥ï¼ˆæ ¸å¿ƒï¼‰ï¼š
        - split('\n')ï¼šä¸¥æ ¼æŒ‰æ¢è¡Œç¬¦åˆ†å‰²
        - ä¸ä½¿ç”¨strip()ï¼šä¿ç•™æ¯è¡ŒåŸå§‹å†…å®¹
        - ç©ºè¡Œå¤„ç†ï¼šåˆ›å»ºç©ºæ®µè½ä¿æŒç‰ˆå¼
        - æ®µè½ç‹¬ç«‹ï¼šæ¯è¡Œä¸€ä¸ªæ®µè½ç¡®ä¿æ¢è¡Œæ•ˆæœ
        
        ä¸ä¹‹å‰ç‰ˆæœ¬åŒºåˆ«ï¼š
        - æ—§ç‰ˆï¼šå¤æ‚çš„æ®µè½å’Œæ¢è¡Œå¤„ç†ï¼Œå®¹æ˜“ä¸¢å¤±æ ¼å¼
        - æ–°ç‰ˆï¼šç®€å•çš„è¡Œçº§å¤„ç†ï¼Œå®Œç¾ä¿ç•™åŸå§‹æ ¼å¼
        
        ä½¿ç”¨åœºæ™¯ï¼š
        - VIPç« èŠ‚ï¼šä¿ç•™ä»åå°è·å–çš„åŸå§‹æ ¼å¼
        - å…è´¹ç« èŠ‚ï¼šä¿ç•™ä»é¡µé¢è§£æçš„æ ¼å¼
        - ä½œè€…æœ‰è¯è¯´ï¼šä¿ç•™ç‰¹æ®Šæ ¼å¼å’Œæ¢è¡Œ
        """
        # åˆ†ç¦»æ­£æ–‡å’Œä½œè€…æœ‰è¯è¯´
        main_text = ""
        author_notes = ""
        
        if 'ã€ä½œè€…æœ‰è¯è¯´ã€‘' in content:
            parts = content.split('ã€ä½œè€…æœ‰è¯è¯´ã€‘', 1)
            main_text = parts[0].strip()
            if len(parts) > 1:
                author_notes = parts[1].strip()
        else:
            main_text = content.strip()
        
        # æ·»åŠ æ­£æ–‡å†…å®¹
        if main_text:
            # æŒ‰è¡Œåˆ†å‰²ï¼Œä¿ç•™æ‰€æœ‰æ¢è¡Œç¬¦å’Œç©ºè¡Œ
            lines = main_text.split('\n')
            for line in lines:
                # ä¿ç•™åŸå§‹å†…å®¹ï¼ŒåŒ…æ‹¬ç©ºè¡Œ
                doc.add_paragraph(line)
        
        # æ·»åŠ ä½œè€…æœ‰è¯è¯´éƒ¨åˆ†
        if author_notes:
            author_heading = doc.add_heading('ä½œè€…æœ‰è¯è¯´', level=2)
            author_heading.runs[0].font.color.rgb = RGBColor(0, 0, 255)
            
            # æŒ‰è¡Œåˆ†å‰²ï¼Œä¿ç•™æ‰€æœ‰æ¢è¡Œç¬¦å’Œç©ºè¡Œ
            lines = author_notes.split('\n')
            for line in lines:
                # ä¿ç•™åŸå§‹å†…å®¹ï¼ŒåŒ…æ‹¬ç©ºè¡Œ
                doc.add_paragraph(line)
    
    def select_novels_to_backup(self, novels):
        """ç”¨æˆ·é€‰æ‹©è¦å¤‡ä»½çš„ä½œå“"""
        if not novels:
            return []
        
        print("\n" + "="*50)
        print("å‘ç°ä»¥ä¸‹ä½œå“ï¼š")
        print("="*50)
        
        for idx, novel in enumerate(novels):
            print(f"{idx+1:2d}. {novel['title']}")
            print(f"     ID: {novel['id']} | å­—æ•°: {novel.get('word_count', 'æœªçŸ¥')} | çŠ¶æ€: {novel.get('status', 'æœªçŸ¥')}")
        
        print("\né€‰æ‹©æ–¹å¼ï¼š")
        print("  è¾“å…¥æ•°å­—é€‰æ‹©å•æœ¬ä½œå“ï¼ˆå¦‚ï¼š1ï¼‰")
        print("  è¾“å…¥å¤šä¸ªæ•°å­—é€‰æ‹©å¤šæœ¬ä½œå“ï¼ˆå¦‚ï¼š1,3,5ï¼‰")
        print("  è¾“å…¥ 'all' æˆ– 'a' å¤‡ä»½å…¨éƒ¨ä½œå“")
        print("  è¾“å…¥ 'quit' æˆ– 'q' é€€å‡ºç¨‹åº")
        
        while True:
            try:
                choice = input("\nè¯·è¾“å…¥é€‰æ‹©: ").strip().lower()
                
                if choice in ['quit', 'q']:
                    print("é€€å‡ºç¨‹åº")
                    return []
                
                if choice in ['all', 'a']:
                    print(f"é€‰æ‹©å¤‡ä»½å…¨éƒ¨ {len(novels)} éƒ¨ä½œå“")
                    return novels
                
                # è§£ææ•°å­—é€‰æ‹©
                selected_indices = []
                for item in choice.split(','):
                    item = item.strip()
                    if item.isdigit():
                        idx = int(item) - 1
                        if 0 <= idx < len(novels):
                            selected_indices.append(idx)
                        else:
                            print(f"æ•°å­— {item} è¶…å‡ºèŒƒå›´ï¼Œè¯·é‡æ–°è¾“å…¥")
                            break
                    else:
                        print(f"æ— æ•ˆè¾“å…¥ '{item}'ï¼Œè¯·é‡æ–°è¾“å…¥")
                        break
                else:
                    if selected_indices:
                        selected_novels = [novels[i] for i in selected_indices]
                        print(f"é€‰æ‹©å¤‡ä»½ {len(selected_novels)} éƒ¨ä½œå“:")
                        for novel in selected_novels:
                            print(f"  - {novel['title']}")
                        return selected_novels
                    else:
                        print("æœªé€‰æ‹©ä»»ä½•ä½œå“ï¼Œè¯·é‡æ–°è¾“å…¥")
                        
            except KeyboardInterrupt:
                print("\n\nç”¨æˆ·ä¸­æ–­ï¼Œé€€å‡ºç¨‹åº")
                return []
            except Exception as e:
                print(f"è¾“å…¥é”™è¯¯: {e}ï¼Œè¯·é‡æ–°è¾“å…¥")
    
    def backup_all_novels(self):
        """å¤‡ä»½ä½œå“ä¸»æµç¨‹"""
        print("æ­£åœ¨åˆå§‹åŒ–...")
        
        # æ£€æŸ¥ç™»å½•çŠ¶æ€
        self.check_login()
        
        # è·å–ä½œå“åˆ—è¡¨
        print("æ­£åœ¨è·å–ä½œå“åˆ—è¡¨...")
        novels = self.get_novel_list()
        
        if not novels:
            print("âŒ æ²¡æœ‰æ‰¾åˆ°ä½œå“")
            return
        
        print(f"âœ“ æˆåŠŸè·å– {len(novels)} éƒ¨ä½œå“")
        
        # ç”¨æˆ·é€‰æ‹©è¦å¤‡ä»½çš„ä½œå“
        selected_novels = self.select_novels_to_backup(novels)
        if not selected_novels:
            return
        
        # ä¿å­˜ä½œå“åˆ—è¡¨ä¿¡æ¯
        with open(os.path.join(self.output_dir, "ä½œå“åˆ—è¡¨.json"), "w", encoding="utf-8") as f:
            json.dump(selected_novels, f, ensure_ascii=False, indent=2)
        
        total_novels = len(selected_novels)
        print(f"\n{'='*50}")
        print(f"å¼€å§‹å¤‡ä»½ {total_novels} éƒ¨ä½œå“")
        print(f"{'='*50}")
        
        # å¤‡ä»½æ¯éƒ¨ä½œå“
        for idx, novel in enumerate(selected_novels):
            print(f"\nâ–¶ [{idx+1}/{total_novels}] å¼€å§‹å¤‡ä»½: {novel['title']}")
            
            # è·å–ç« èŠ‚åˆ—è¡¨
            chapters = self.get_chapters(novel['link'])
            
            if chapters:
                # åˆ›å»ºDOCXæ–‡ä»¶
                self.create_docx_with_realtime_save(novel, chapters)
            else:
                print(f"âŒ æœªæ‰¾åˆ°ç« èŠ‚ï¼Œè·³è¿‡: {novel['title']}")
            
            # ä½œå“é—´å»¶è¿Ÿ
            if idx < total_novels - 1:
                delay = random.uniform(2.0, 4.0)
                print(f"ç­‰å¾… {delay:.1f} ç§’åç»§ç»­...")
                time.sleep(delay)
        
        print(f"\n{'='*50}")
        print(f"ğŸ‰ å¤‡ä»½å®Œæˆï¼æ–‡ä»¶å·²ä¿å­˜åˆ°: {self.output_dir}")
        print(f"{'='*50}")


if __name__ == "__main__":
    print("""
    â•”â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•—
    â•‘                  æ™‹æ±Ÿæ–‡å­¦åŸä½œå“å¤‡ä»½å·¥å…· v5.0                   â•‘
    â•‘                     (ä¼˜åŒ–ç‰ˆ - æ”¯æŒé€‰æ‹©å¤‡ä»½)                    â•‘
    â• â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•£
    â•‘  åŠŸèƒ½ç‰¹æ€§:                                                     â•‘
    â•‘  â€¢ ä» my_cookie.txt æ–‡ä»¶è¯»å–Cookie                            â•‘
    â•‘  â€¢ æ”¯æŒé€‰æ‹©å•æœ¬æˆ–å¤šæœ¬ä½œå“å¤‡ä»½                                  â•‘
    â•‘  â€¢ å®æ—¶ä¿å­˜ï¼Œè¾¹ä¸‹è½½è¾¹ç”ŸæˆDOCXæ–‡ä»¶                             â•‘
    â•‘  â€¢ å®Œæ•´ä¿ç•™æ­£æ–‡æ ¼å¼å’Œä½œè€…æœ‰è¯è¯´                               â•‘
    â•‘  â€¢ è‡ªåŠ¨æ·»åŠ ç« èŠ‚ç¼–å·å’Œå±‚çº§æ ‡é¢˜                                 â•‘
    â•‘                                                                â•‘
    â•‘  ä½¿ç”¨æ–¹æ³•:                                                     â•‘
    â•‘  1. å‡†å¤‡Cookie: ç™»å½•æ™‹æ±Ÿâ†’F12â†’Networkâ†’å¤åˆ¶Cookieåˆ°txtæ–‡ä»¶     â•‘
    â•‘  2. è¿è¡Œç¨‹åº: python jjwxc_col.py                            â•‘
    â•‘  3. é€‰æ‹©ä½œå“: æ ¹æ®æç¤ºé€‰æ‹©è¦å¤‡ä»½çš„ä½œå“                        â•‘
    â•‘  4. æŸ¥çœ‹ç»“æœ: å¤‡ä»½å®ŒæˆåæŸ¥çœ‹ç”Ÿæˆçš„DOCXæ–‡ä»¶                   â•‘
    â•šâ•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    """)
    
    # æ£€æŸ¥Cookieæ–‡ä»¶
    if not os.path.exists(COOKIE_FILE):
        print(f"âŒ æœªæ‰¾åˆ° {COOKIE_FILE} æ–‡ä»¶")
        print("\nğŸ“ Cookieè·å–æ­¥éª¤:")
        print("1. ä½¿ç”¨æµè§ˆå™¨ç™»å½•æ™‹æ±Ÿæ–‡å­¦åŸä½œè€…åå°")
        print("2. æŒ‰F12æ‰“å¼€å¼€å‘è€…å·¥å…·")
        print("3. åˆ‡æ¢åˆ°Network(ç½‘ç»œ)é€‰é¡¹å¡")
        print("4. åˆ·æ–°é¡µé¢ï¼Œç‚¹å‡»ä»»æ„è¯·æ±‚")
        print("5. åœ¨Request Headersä¸­æ‰¾åˆ°'Cookie'å­—æ®µ")
        print("6. å¤åˆ¶å®Œæ•´çš„Cookieå€¼")
        print(f"7. åˆ›å»º {COOKIE_FILE} æ–‡ä»¶ï¼Œç²˜è´´Cookieå†…å®¹å¹¶ä¿å­˜")
        print("\næŒ‰å›è½¦é”®é€€å‡º...")
        input()
        exit(1)
    
    # å¯åŠ¨å¤‡ä»½å·¥å…·
    try:
        tool = JJWXCBackupTool()
        tool.backup_all_novels()
    except KeyboardInterrupt:
        print("\n\nç”¨æˆ·ä¸­æ–­ç¨‹åº")
    except Exception as e:
        print(f"\nç¨‹åºè¿è¡Œå‡ºé”™: {e}")
        print("è¯·æ£€æŸ¥ç½‘ç»œè¿æ¥å’ŒCookieæ˜¯å¦æœ‰æ•ˆ")
    finally:
        print("\nç¨‹åºç»“æŸï¼ŒæŒ‰å›è½¦é”®é€€å‡º...")
        input()