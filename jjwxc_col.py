#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
=================================================================
                晋江文学城作品备份工具 v5.0
=================================================================
功能说明:
- 通过作者后台批量备份作品到DOCX格式
- 支持免费章节和VIP章节的内容获取  
- 自动保留原文格式和作者有话说
- 实时保存，边下载边生成文档

技术特性:
- Cookie认证：支持复杂Cookie格式解析
- 格式保留：完整保持换行符和空行
- 智能重试：网络异常自动重试机制

使用前提:
1. 需要晋江作者账号的登录Cookie
2. 在my_cookie.txt文件中粘贴完整Cookie字符串
3. 确保网络连接稳定

=================================================================
"""

import os
import time
import random
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import json
from datetime import datetime
import urllib.parse

# ========================= 全局配置 =========================
COOKIE_FILE = "my_cookie.txt"  # Cookie文件路径

class JJWXCBackupTool:
    """
    晋江文学城作品备份工具主类
    
    主要功能：
    1. Cookie认证和会话管理
    2. 作品列表获取和解析
    3. 章节内容抓取（免费+VIP）
    4. DOCX文档生成和格式化
    
    使用流程：
    init() -> check_login() -> get_novel_list() -> select_novels() -> backup_novels()
    """
    
    def __init__(self):
        """
        初始化备份工具
        
        功能：
        - 创建输出目录结构 (backup/YYYYMMDD_HHMMSS/)
        - 初始化HTTP会话和请求头
        - 加载Cookie文件并解析认证信息
        - 配置网络重试策略
        """
        # 创建输出目录 - 使用timestamp确保唯一性
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        self.output_dir = os.path.join("backup", timestamp)
        os.makedirs(self.output_dir, exist_ok=True)
        print(f"输出目录: {self.output_dir}")
        
        # 设置HTTP会话 - 保持Cookie和连接复用
        self.session = requests.Session()
        self.headers = self.get_default_headers()
        
        # 初始化作者后台URL
        self.author_backend_url = None
        
        # 加载并解析Cookie文件
        cookie_count = self.load_cookie()
        print(f"已设置 {cookie_count} 个Cookie参数")
        
        # 设置请求重试策略 - 应对网络波动
        self.session.mount('https://', requests.adapters.HTTPAdapter(
            max_retries=3,
            pool_connections=10,
            pool_maxsize=20
        ))

    def get_default_headers(self):
        """
        获取默认HTTP请求头
        
        返回：
            dict: 包含User-Agent、Accept等标准浏览器请求头的字典
            
        用途：
            模拟真实浏览器访问，避免被网站反爬虫策略拦截
        """
        return {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
            'Connection': 'keep-alive',
            'Referer': 'https://www.jjwxc.net/',
            'Cache-Control': 'no-cache',
            'Pragma': 'no-cache'
        }
    
    def decode_unicode_escape(self, s):
        """
        解码Unicode转义序列
        
        参数：
            s (str): 包含%uXXXX格式Unicode转义的字符串
            
        返回：
            str: 解码后的字符串
            
        说明：
            晋江Cookie中可能包含%u格式的Unicode转义字符
            需要特殊处理才能正确解析中文字符
        """
        def replace_unicode(match):
            return chr(int(match.group(1), 16))
        return re.sub(r'%u([0-9a-fA-F]{4})', replace_unicode, s)
    
    def load_cookie(self):
        """
        从文件加载并解析Cookie
        
        返回：
            int: 成功解析的Cookie数量
            
        功能：
        1. 读取my_cookie.txt文件内容
        2. 处理Unicode转义序列(%uXXXX)
        3. 智能解析复杂Cookie格式（包括JSON值）
        4. 设置到HTTP会话中
        
        Cookie格式支持：
        - 标准键值对: key=value; key2=value2
        - JSON值: key={"json":"value"}  
        - URL编码值: key=%E4%B8%AD%E6%96%87
        - Unicode转义: key=%u4E2D%u6587
        
        错误处理：
        - 文件不存在：返回0
        - 解析失败：跳过该Cookie并继续
        - JSON验证：确保JSON格式Cookie的有效性
        """
        cookie_count = 0
        if os.path.exists(COOKIE_FILE):
            try:
                with open(COOKIE_FILE, 'r', encoding='utf-8') as f:
                    # 读取原始Cookie内容
                    raw_cookie = f.read().strip()
                    print(f"原始Cookie内容: {raw_cookie[:100]}...")
                    
                    # 第一步：处理Unicode转义序列 (%uXXXX)
                    decoded_cookie = self.decode_unicode_escape(raw_cookie)
                    
                    # 第二步：智能Cookie解析 - 处理包含JSON的复杂Cookie
                    cookies_dict = {}
                    current_pos = 0
                    
                    while current_pos < len(decoded_cookie):
                        # 找到下一个等号
                        eq_pos = decoded_cookie.find('=', current_pos)
                        if eq_pos == -1:
                            break
                            
                        # 提取键名（去掉前面的分号和空格）
                        key_start = current_pos
                        if decoded_cookie[current_pos:current_pos+1] == ';':
                            key_start += 1
                        key = decoded_cookie[key_start:eq_pos].strip()
                        
                        # 找值的部分
                        value_start = eq_pos + 1
                        
                        # 如果值以 %7B 或 { 开头，可能是JSON，需要特殊处理
                        if (decoded_cookie[value_start:value_start+3] == '%7B' or 
                            decoded_cookie[value_start:value_start+1] == '{'):
                            # JSON值：找到匹配的结束位置
                            brace_count = 0
                            value_end = value_start
                            in_string = False
                            escape_next = False
                            
                            for i in range(value_start, len(decoded_cookie)):
                                char = decoded_cookie[i]
                                
                                if escape_next:
                                    escape_next = False
                                    continue
                                    
                                if char == '\\':
                                    escape_next = True
                                    continue
                                    
                                if char == '"' and not escape_next:
                                    in_string = not in_string
                                    continue
                                    
                                if not in_string:
                                    if char == '{' or decoded_cookie[i:i+3] == '%7B':
                                        brace_count += 1
                                        if decoded_cookie[i:i+3] == '%7B':
                                            i += 2  # 跳过 %7B 的其余部分
                                    elif char == '}' or decoded_cookie[i:i+3] == '%7D':
                                        brace_count -= 1
                                        if decoded_cookie[i:i+3] == '%7D':
                                            i += 2
                                        if brace_count == 0:
                                            value_end = i + 1
                                            break
                                    elif char == ';' and brace_count == 0:
                                        value_end = i
                                        break
                                        
                                value_end = i + 1
                        else:
                            # 普通值：找到下一个分号或字符串结尾
                            value_end = decoded_cookie.find(';', value_start)
                            if value_end == -1:
                                value_end = len(decoded_cookie)
                        
                        # 提取值
                        value = decoded_cookie[value_start:value_end].strip()
                        
                        # 处理值的URL解码
                        if key and value:
                            try:
                                # 对于JSON格式的Cookie值，谨慎处理URL解码
                                if value.startswith('%7B') or value.startswith('{'):
                                    # 尝试解码，但如果失败就保持原样
                                    try:
                                        decoded_value = urllib.parse.unquote(value)
                                        # 验证解码后的JSON是否有效
                                        if decoded_value.startswith('{') and decoded_value.endswith('}'):
                                            json.loads(decoded_value)  # 验证JSON格式
                                    except:
                                        decoded_value = value  # 解码失败，保持原样
                                else:
                                    # 普通值，直接解码
                                    decoded_value = urllib.parse.unquote(value)
                                
                                cookies_dict[key] = decoded_value
                                self.session.cookies.set(key, decoded_value)
                                print(f"  → 设置Cookie: {key}={decoded_value[:50]}...")
                                cookie_count += 1
                                
                            except Exception as decode_error:
                                print(f"  × 跳过无效Cookie: {key} (解码错误: {decode_error})")
                        
                        # 移动到下一个Cookie
                        current_pos = value_end
                        if current_pos < len(decoded_cookie) and decoded_cookie[current_pos] == ';':
                            current_pos += 1
                    
                    print(f"成功解析 {cookie_count} 个Cookie")
                    return cookie_count
                    
            except Exception as e:
                print(f"Cookie文件解析错误: {e}")
        
        return 0
    
    def check_login(self):
        """检查登录状态"""
        self.author_backend_url = "https://my.jjwxc.net/backend/oneauthor_login.php"
        try:
            print("正在检查登录状态...")
            response = self.session.get(self.author_backend_url, headers=self.headers, timeout=15)
            response.encoding = 'gb18030'
            html = response.text

            # 判断页面是否包含登录提示
            if "晋江文学城" in html:
                print("登入成功")
                return True
            elif "请登录" in html or "登录晋江作者后台" in html or "账号" in html:
                print("未登录晋江作者后台，请检查Cookie是否有效")
                return False
        except Exception as e:
            print(f"检查登录状态时出错: {e}")
            return False
        
    def get_novel_list(self):
        """获取作者作品列表"""
        author_url = "https://my.jjwxc.net/backend/oneauthor_login.php"
        
        try:
            print(f"获取作品列表: {author_url}")
            response = self.session.get(author_url, headers=self.headers, timeout=20)
            
            # 设置正确的编码
            response.encoding = 'gb18030'  # 晋江使用gb18030编码
            
            soup = BeautifulSoup(response.content, 'html.parser', from_encoding='gb18030')
            novels = []
            
            # 保存页面用于调试（已禁用，如需调试请取消注释）
            # with open(os.path.join(self.output_dir, "novel_list.html"), "w", encoding="utf-8") as f:
            #     f.write(response.text)
            # print("作品列表页面已保存: novel_list.html")
            
            # 查找作品管理链接
            # 在晋江后台，作品管理链接的格式是: managenovel.php?novelid=XXXXX
            novel_links = soup.find_all('a', href=lambda x: x and 'managenovel.php?novelid=' in x)
            
            if novel_links:
                print(f"找到 {len(novel_links)} 个作品管理链接")
                
                for link in novel_links:
                    # 提取作品ID
                    href = link['href']
                    novel_id_match = re.search(r'novelid=(\d+)', href)
                    if novel_id_match:
                        novel_id = novel_id_match.group(1)
                        
                        # 查找作品标题（在同一行的其他位置）
                        row = link.find_parent('tr')
                        if row:
                            # 查找作品标题链接（指向onebook.php的链接）
                            title_link = row.find('a', href=lambda x: x and f'onebook.php?novelid={novel_id}' in x)
                            if title_link:
                                title = title_link.get_text(strip=True)
                                
                        # 提取其他信息
                        cells = row.find_all('td')
                        if len(cells) >= 10:
                            try:
                                # 根据HTML结构提取信息
                                category = cells[2].get_text(strip=True) if len(cells) > 2 else "未知"
                                subcategory = cells[3].get_text(strip=True) if len(cells) > 3 else "未知"
                                chapter_count = cells[5].get_text(strip=True) if len(cells) > 5 else "0"
                                word_count = cells[6].get_text(strip=True) if len(cells) > 6 else "0"
                                status = cells[12].get_text(strip=True) if len(cells) > 12 else "未知"
                                
                                novels.append({
                                    'id': novel_id,
                                    'title': title,
                                    'link': href,
                                    'view_link': title_link['href'],
                                    'status': status,
                                    'word_count': word_count,
                                    'chapter_count': chapter_count,
                                    'category': f"{category}-{subcategory}"
                                })
                                
                            except Exception as e:
                                print(f"解析作品信息出错 {novel_id}: {e}")
                                novels.append({
                                    'id': novel_id,
                                    'title': title,
                                    'link': href,
                                    'view_link': title_link['href'],
                                    'status': "未知",
                                    'word_count': "未知",
                                    'chapter_count': "未知",
                                    'category': "未知"
                                })
                
                print(f"成功解析 {len(novels)} 部作品")
                return novels
            else:
                print("未找到作品管理链接")
                # 备用方法：查找onebook.php链接
                onebook_links = soup.find_all('a', href=lambda x: x and 'onebook.php?novelid=' in x)
                if onebook_links:
                    print(f"找到 {len(onebook_links)} 个作品阅读链接")
                    for idx, link in enumerate(onebook_links):
                        href = link['href']
                        novel_id_match = re.search(r'novelid=(\d+)', href)
                        if novel_id_match:
                            novel_id = novel_id_match.group(1)
                            title = link.get_text(strip=True)
                            
                            novels.append({
                                'id': novel_id,
                                'title': title,
                                'link': f"//my.jjwxc.net/backend/managenovel.php?novelid={novel_id}",
                                'view_link': href,
                                'status': "未知",
                                'word_count': "未知",
                                'chapter_count': "未知",
                                'category': "未知"
                            })
                    
                    return novels
                else:
                    print("也未找到作品阅读链接")
                    return []
            
        except Exception as e:
            print(f"获取作品列表出错: {str(e)}")
            return []
    
    def get_intro_from_backend(self, novel_id):
        """从作者后台获取作品简介"""
        try:
            backend_url = f"https://my.jjwxc.net/backend/managenovel.php?novelid={novel_id}"
            print(f"访问后台章节管理页面: {backend_url}")
            headers = self.headers.copy()
            headers['Referer'] = 'https://my.jjwxc.net/backend/'
            response = self.session.get(backend_url, headers=headers, timeout=30)
            response.encoding = 'gb18030'
            
            soup = BeautifulSoup(response.content, 'html.parser', from_encoding='gb18030')
            novel_intro = ""
            intro_textarea = soup.find('textarea', {'id': 'novelintro'})
            if intro_textarea:
                novel_intro = intro_textarea.get_text(strip=True)
                print(f"获取到作品简介: {len(novel_intro)} 字符")
            return novel_intro
        except Exception as e:
            print(f"获取作品简介失败: {e}")
            return ""

    def get_chapters(self, novel_link):
        """
        获取作品的完整章节列表
        
        参数：
            novel_link (str): 作品管理页面链接
            
        返回：
            list: 章节信息列表，每个元素包含：
                - id: 章节ID
                - title: 章节标题  
                - link: 章节访问链接
                - chapter_number: 章节编号
                - is_vip: 是否VIP章节
                
        功能流程：
        1. 从链接提取作品ID
        2. 访问后台章节管理页面
        3. 解析章节表格（包括隐藏行）
        4. 区分免费和VIP章节
        5. 按章节编号排序
        
        章节识别逻辑：
        - VIP章节：链接包含'onebook_vip.php'
        - 免费章节：其他章节
        - 隐藏章节：包括display:none的行
        """
        if not novel_link:
            return []
            
        try:
            # 提取novelid
            novel_id_match = re.search(r'novelid=(\d+)', novel_link)
            if not novel_id_match:
                print(f"无法从链接中提取作品ID: {novel_link}")
                return []
            novel_id = novel_id_match.group(1)
            backend_url = f"https://my.jjwxc.net/backend/managenovel.php?novelid={novel_id}"
            print(f"获取所有章节列表: {backend_url}")
            response = self.session.get(backend_url, headers=self.headers, timeout=30)
            response.encoding = 'gb18030'
            
            soup = BeautifulSoup(response.content, 'html.parser', from_encoding='gb18030')
            chapters = []
            
            # 查找章节表格，包括隐藏的章节行
            # 查找所有包含章节信息的tr行，包括style="display:none;"的隐藏行
            chapter_rows = soup.find_all('tr', {'bgcolor': '#eefaee'})
            
            for tr in chapter_rows:
                # 查找章节ID input
                chapter_id_input = tr.find('input', {'name': 'chapterid'})
                if not chapter_id_input:
                    continue
                    
                chapter_id = chapter_id_input.get('value')
                if not chapter_id:
                    continue
                
                # 查找章节标题链接
                title_link = tr.find('a', href=True)
                if not title_link:
                    continue
                    
                title = title_link.get_text(strip=True)
                href = title_link.get('href')
                
                # 判断是否VIP章节 - 通过链接或标题中的VIP标识
                is_vip = False
                if 'onebook_vip.php' in href or '[VIP]' in title_link.get_text():
                    is_vip = True
                
                # 提取章节编号（从第二个td中获取）
                chapter_num_td = tr.find_all('td')[1]  # 第二个td包含章节编号
                chapter_number = int(chapter_num_td.get_text(strip=True))
                
                # 构建完整的访问链接
                if not href.startswith('http'):
                    if href.startswith('//'):
                        link = f"https:{href}"
                    else:
                        link = f"https://www.jjwxc.net{href}"
                else:
                    link = href
                
                chapters.append({
                    'id': chapter_id,
                    'title': title,
                    'link': link,
                    'chapter_number': chapter_number,
                    'is_vip': is_vip
                })
            
            # 按章节编号排序
            chapters.sort(key=lambda x: x['chapter_number'])
            
            vip_count = sum(1 for c in chapters if c['is_vip'])
            free_count = len(chapters) - vip_count
            print(f"成功解析 {len(chapters)} 个章节，其中免费章节数量：{free_count}，VIP章节数量：{vip_count}")
            return chapters
        except Exception as e:
            print(f"获取章节列表出错: {str(e)}")
            return []
    
    def get_chapter_content(self, chapter_link, is_vip=False):
        """
        获取章节内容（支持免费和VIP章节）
        
        参数：
            chapter_link (str): 章节链接
            is_vip (bool): 是否为VIP章节
            
        返回：
            str: 章节完整内容（包含正文和作者有话说）
            
        免费章节处理：
        1. 直接访问章节页面
        2. 解析novelbody div中的正文内容
        3. 解析note_str div中的作者有话说
        4. 保留换行格式和HTML结构
        
        VIP章节处理（核心功能）：
        1. 从章节链接提取novelid和chapterid
        2. 构建后台编辑页面URL (chaptermodify.php)
        3. 从textarea获取未加密的原始内容：
           - name='content': 章节正文
           - name='note': 作者有话说
        4. 完整保留原始格式（换行、空行、特殊字符）
        5. 清理HTML实体编码但保持文本结构
        
        格式保留策略：
        - 使用textarea.string获取原始文本
        - 保留所有\n换行符和空行
        - 只处理HTML实体转义（&lt; &gt; &amp;等）
        - 不做额外的文本清理或格式化
        
        错误处理：
        - 链接无效：返回错误信息
        - 网络异常：返回异常描述
        - 内容为空：返回获取失败提示
        """
        if not chapter_link:
            return "章节链接无效"
        try:
            # 免费章节处理逻辑
            if not is_vip:
                if not chapter_link.startswith('http'):
                    chapter_link = f"https://www.jjwxc.net{chapter_link}"
                print(f"  获取免费章节内容...")
                response = self.session.get(chapter_link, headers=self.headers, timeout=30)
                response.encoding = 'gb18030'
                soup = BeautifulSoup(response.content, 'html.parser', from_encoding='gb18030')
                main_content = ""
                novelbody_div = soup.find('div', class_='novelbody')
                if novelbody_div:
                    text_container = novelbody_div.select_one('div > div')
                    if text_container:
                        content_parts = []
                        for element in text_container.children:
                            if getattr(element, 'name', None) == 'br':
                                content_parts.append('\n')
                            elif getattr(element, 'string', None) and element.string.strip():
                                content_parts.append(element.string.strip())
                        main_content = ''.join(content_parts)
                author_notes = ""
                note_wrapper = soup.find('div', id='note_danmu_wrapper')
                if note_wrapper:
                    note_str = note_wrapper.find('div', id='note_str')
                    if note_str:
                        html_content = str(note_str)
                        html_content = re.sub(r'<br\s*/?>', '\n', html_content, flags=re.IGNORECASE)
                        clean_soup = BeautifulSoup(html_content, 'html.parser')
                        author_notes = clean_soup.get_text(strip=True)
                result_parts = []
                if main_content and len(main_content) > 20:
                    result_parts.append(main_content)
                if author_notes and len(author_notes) > 10:
                    result_parts.append('\n\n【作者有话说】')
                    result_parts.append(author_notes)
                if result_parts:
                    result = ''.join(result_parts)
                    if len(result.strip()) > 30:
                        return result
                return "内容获取失败：未找到有效内容"
            # VIP章节处理逻辑 - 从作者后台编辑页面获取未加密内容
            else:
                # 从章节链接中提取novelid和chapterid
                novel_id_match = re.search(r'novelid=(\d+)', chapter_link)
                chapter_id_match = re.search(r'chapterid=(\d+)', chapter_link)
                
                if not novel_id_match or not chapter_id_match:
                    return "VIP章节链接格式错误"
                
                novel_id = novel_id_match.group(1)
                chapter_id = chapter_id_match.group(1)
                
                # 构建后台编辑页面链接
                edit_url = f"https://my.jjwxc.net/backend/chaptermodify.php?novelid={novel_id}&chapterid={chapter_id}"
                print(f"  获取VIP章节内容（从编辑页面）...")
                
                headers = self.headers.copy()
                headers['Referer'] = f'https://my.jjwxc.net/backend/managenovel.php?novelid={novel_id}'
                
                response = self.session.get(edit_url, headers=headers, timeout=30)
                response.encoding = 'gb18030'
                soup = BeautifulSoup(response.content, 'html.parser', from_encoding='gb18030')
                
                main_content = ""
                author_notes = ""
                
                # 从编辑页面的textarea获取正文内容
                chapterbody_textarea = soup.find('textarea', {'name': 'content'})
                if chapterbody_textarea:
                    # 获取原始文本内容，保留所有格式
                    main_content = chapterbody_textarea.string or chapterbody_textarea.get_text()
                    # 如果没有内容，尝试从textarea内部获取
                    if not main_content.strip():
                        main_content = ''.join(str(content) for content in chapterbody_textarea.contents)
                    
                    # 只清理HTML实体编码，保留所有换行和空行
                    main_content = main_content.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
                    main_content = main_content.replace('&quot;', '"').replace('&#039;', "'")
                    main_content = main_content.replace('&nbsp;', ' ')  # 处理非断行空格
                
                # 获取作者有话说
                authornote_textarea = soup.find('textarea', {'name': 'note'})
                if authornote_textarea:
                    # 获取原始文本内容，保留所有格式
                    author_notes = authornote_textarea.string or authornote_textarea.get_text()
                    # 如果没有内容，尝试从textarea内部获取
                    if not author_notes.strip():
                        author_notes = ''.join(str(content) for content in authornote_textarea.contents)
                    
                    # 只清理HTML实体编码，保留所有换行和空行
                    author_notes = author_notes.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
                    author_notes = author_notes.replace('&quot;', '"').replace('&#039;', "'")
                    author_notes = author_notes.replace('&nbsp;', ' ')  # 处理非断行空格
                
                result_parts = []
                if main_content and len(main_content.strip()) > 20:
                    result_parts.append(main_content)
                if author_notes and len(author_notes.strip()) > 10:
                    result_parts.append('\n\n【作者有话说】\n')
                    result_parts.append(author_notes)
                if result_parts:
                    result = ''.join(result_parts)
                    if len(result.strip()) > 30:
                        return result
                return "VIP内容获取失败：未找到有效内容"
        except Exception as e:
            print(f"  章节内容获取出错: {str(e)}")
            return f"内容获取失败：{str(e)}"

    def create_docx_with_realtime_save(self, novel, chapters):
        """
        创建DOCX文档并实时保存章节内容
        
        参数：
            novel (dict): 作品信息字典
            chapters (list): 章节列表
            
        功能特性：
        1. 文档结构创建：
           - 作品标题（0级标题，居中）
           - 作品信息（ID、字数、状态）
           - 作品简介（从后台获取）
           - 分页符分隔
           
        2. 章节处理：
           - 逐章节获取和添加内容
           - 实时保存（每章节保存一次）
           - 章节标题格式化（第X章 标题）
           - 章节间分隔符
           
        3. 内容格式化：
           - 调用_add_content_to_doc处理正文
           - 自动分离作者有话说
           - 保留原始换行和空行
           - 错误章节标红显示
           
        4. 实时保存机制：
           - 创建初始文档结构立即保存
           - 每添加一章节内容后保存
           - 用户可随时打开查看进度
           - 避免程序中断导致数据丢失
           
        5. 文件命名：
           - 清理标题中的非法字符
           - 生成safe的文件名
           - 保存到backup/timestamp/目录
           
        6. 错误处理：
           - 章节获取失败：标红显示错误信息
           - 网络异常：继续处理下一章节
           - 文档保存异常：记录错误并继续
           
        7. 进度显示：
           - 显示当前章节进度 [X/总数]
           - 预估剩余时间
           - 章节获取状态反馈
        """
        if not chapters:
            print(f"没有找到章节内容，跳过 {novel['title']}")
            return
        
        try:
            # 创建Word文档
            doc = Document()
            
            # 添加作品标题（最高级标题）
            title_paragraph = doc.add_heading(novel['title'], level=0)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加作品基本信息
            info_paragraph = doc.add_paragraph()
            info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            info_run = info_paragraph.add_run(
                f"作品ID: {novel['id']} | "
                f"字数: {novel.get('word_count', '未知')} | "
                f"状态: {novel.get('status', '未知')}"
            )
            info_run.font.size = Pt(10)
            
            # 获取作品简介并插入到状态下方
            novel_intro = self.get_intro_from_backend(novel['id'])
            if novel_intro:
                intro_paragraph = doc.add_paragraph(novel_intro)
                intro_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                intro_paragraph.runs[0].font.size = Pt(11)
            
            # 添加分页符
            doc.add_page_break()
            
            # 准备文件名和路径
            filename = self._clean_filename(novel['title'])
            filepath = os.path.join(self.output_dir, f"{filename}.docx")
            
            total_chapters = len(chapters)
            print(f"开始处理: {novel['title']} ({total_chapters}章)")
            print(f"文档将保存为: {filepath}")
            
            # 先保存初始文档结构
            doc.save(filepath)
            print(f"✓ 已创建初始文档，可以打开查看")
            
            # 逐章节处理并实时保存
            for idx, chapter in enumerate(chapters):
                try:
                    # 添加章节标题（带章节编号）
                    chapter_title = f"第{chapter.get('chapter_number', idx+1)}章 {chapter['title']}"
                    doc.add_heading(chapter_title, level=1)
                    
                    # 获取章节内容
                    print(f"正在获取: {chapter_title} [{idx+1}/{total_chapters}]")
                    content = self.get_chapter_content(chapter['link'], chapter.get('is_vip', False))
                    
                    # 检查内容是否有效
                    if content and not content.startswith("内容获取失败") and not content.startswith("章节链接无效"):
                        self._add_content_to_doc(doc, content)
                    else:
                        # 内容获取失败的情况
                        error_paragraph = doc.add_paragraph(f"[章节内容获取失败: {content}]")
                        error_paragraph.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    
                    # 添加章节分隔符
                    if idx < total_chapters - 1:
                        doc.add_paragraph()
                        separator = doc.add_paragraph("─" * 50)
                        separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        doc.add_paragraph()
                    
                    # 实时保存文档
                    doc.save(filepath)
                    print(f"✓ 已保存 [{idx+1}/{total_chapters}]")
                    
                except Exception as e:
                    print(f"处理章节出错: {str(e)}")
                    error_paragraph = doc.add_paragraph(f"[章节处理错误: {chapter['title']} - {str(e)}]")
                    error_paragraph.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    doc.save(filepath)
                
                # 延迟避免请求过快
                time.sleep(random.uniform(1.0, 2.0))
            
            print(f"✓ 完成保存: {novel['title']}")
            
        except Exception as e:
            print(f"创建文档出错: {str(e)}")
    
    def _clean_filename(self, filename):
        """清理文件名中的非法字符"""
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        if not filename.strip() or filename.strip() == '_':
            filename = f"novel_{int(time.time())}"
        
        return filename
    
    def _add_content_to_doc(self, doc, content):
        """
        将章节内容添加到DOCX文档中（格式保留版）
        
        参数：
            doc: python-docx Document对象
            content (str): 章节完整内容
            
        功能：
        1. 内容分离：
           - 以【作者有话说】为分界点
           - 分离正文和作者有话说两部分
           - 去除首尾空白但保留内部格式
           
        2. 正文处理：
           - 按换行符(\n)分割为行
           - 每行创建独立段落
           - 完整保留空行（空段落）
           - 不做任何文本清理或去空格
           
        3. 作者有话说处理：
           - 添加蓝色二级标题"作者有话说"
           - 同样按行处理内容
           - 保持与正文相同的格式处理方式
           
        格式保留策略（核心）：
        - split('\n')：严格按换行符分割
        - 不使用strip()：保留每行原始内容
        - 空行处理：创建空段落保持版式
        - 段落独立：每行一个段落确保换行效果
        
        与之前版本区别：
        - 旧版：复杂的段落和换行处理，容易丢失格式
        - 新版：简单的行级处理，完美保留原始格式
        
        使用场景：
        - VIP章节：保留从后台获取的原始格式
        - 免费章节：保留从页面解析的格式
        - 作者有话说：保留特殊格式和换行
        """
        # 分离正文和作者有话说
        main_text = ""
        author_notes = ""
        
        if '【作者有话说】' in content:
            parts = content.split('【作者有话说】', 1)
            main_text = parts[0].strip()
            if len(parts) > 1:
                author_notes = parts[1].strip()
        else:
            main_text = content.strip()
        
        # 添加正文内容
        if main_text:
            # 按行分割，保留所有换行符和空行
            lines = main_text.split('\n')
            for line in lines:
                # 保留原始内容，包括空行
                doc.add_paragraph(line)
        
        # 添加作者有话说部分
        if author_notes:
            author_heading = doc.add_heading('作者有话说', level=2)
            author_heading.runs[0].font.color.rgb = RGBColor(0, 0, 255)
            
            # 按行分割，保留所有换行符和空行
            lines = author_notes.split('\n')
            for line in lines:
                # 保留原始内容，包括空行
                doc.add_paragraph(line)
    
    def select_novels_to_backup(self, novels):
        """用户选择要备份的作品"""
        if not novels:
            return []
        
        print("\n" + "="*50)
        print("发现以下作品：")
        print("="*50)
        
        for idx, novel in enumerate(novels):
            print(f"{idx+1:2d}. {novel['title']}")
            print(f"     ID: {novel['id']} | 字数: {novel.get('word_count', '未知')} | 状态: {novel.get('status', '未知')}")
        
        print("\n选择方式：")
        print("  输入数字选择单本作品（如：1）")
        print("  输入多个数字选择多本作品（如：1,3,5）")
        print("  输入 'all' 或 'a' 备份全部作品")
        print("  输入 'quit' 或 'q' 退出程序")
        
        while True:
            try:
                choice = input("\n请输入选择: ").strip().lower()
                
                if choice in ['quit', 'q']:
                    print("退出程序")
                    return []
                
                if choice in ['all', 'a']:
                    print(f"选择备份全部 {len(novels)} 部作品")
                    return novels
                
                # 解析数字选择
                selected_indices = []
                for item in choice.split(','):
                    item = item.strip()
                    if item.isdigit():
                        idx = int(item) - 1
                        if 0 <= idx < len(novels):
                            selected_indices.append(idx)
                        else:
                            print(f"数字 {item} 超出范围，请重新输入")
                            break
                    else:
                        print(f"无效输入 '{item}'，请重新输入")
                        break
                else:
                    if selected_indices:
                        selected_novels = [novels[i] for i in selected_indices]
                        print(f"选择备份 {len(selected_novels)} 部作品:")
                        for novel in selected_novels:
                            print(f"  - {novel['title']}")
                        return selected_novels
                    else:
                        print("未选择任何作品，请重新输入")
                        
            except KeyboardInterrupt:
                print("\n\n用户中断，退出程序")
                return []
            except Exception as e:
                print(f"输入错误: {e}，请重新输入")
    
    def backup_all_novels(self):
        """备份作品主流程"""
        print("正在初始化...")
        
        # 检查登录状态
        self.check_login()
        
        # 获取作品列表
        print("正在获取作品列表...")
        novels = self.get_novel_list()
        
        if not novels:
            print("❌ 没有找到作品")
            return
        
        print(f"✓ 成功获取 {len(novels)} 部作品")
        
        # 用户选择要备份的作品
        selected_novels = self.select_novels_to_backup(novels)
        if not selected_novels:
            return
        
        # 保存作品列表信息
        with open(os.path.join(self.output_dir, "作品列表.json"), "w", encoding="utf-8") as f:
            json.dump(selected_novels, f, ensure_ascii=False, indent=2)
        
        total_novels = len(selected_novels)
        print(f"\n{'='*50}")
        print(f"开始备份 {total_novels} 部作品")
        print(f"{'='*50}")
        
        # 备份每部作品
        for idx, novel in enumerate(selected_novels):
            print(f"\n▶ [{idx+1}/{total_novels}] 开始备份: {novel['title']}")
            
            # 获取章节列表
            chapters = self.get_chapters(novel['link'])
            
            if chapters:
                # 创建DOCX文件
                self.create_docx_with_realtime_save(novel, chapters)
            else:
                print(f"❌ 未找到章节，跳过: {novel['title']}")
            
            # 作品间延迟
            if idx < total_novels - 1:
                delay = random.uniform(2.0, 4.0)
                print(f"等待 {delay:.1f} 秒后继续...")
                time.sleep(delay)
        
        print(f"\n{'='*50}")
        print(f"🎉 备份完成！文件已保存到: {self.output_dir}")
        print(f"{'='*50}")


if __name__ == "__main__":
    print("""
    ╔════════════════════════════════════════════════════════════════╗
    ║                  晋江文学城作品备份工具 v5.0                   ║
    ║                     (优化版 - 支持选择备份)                    ║
    ╠════════════════════════════════════════════════════════════════╣
    ║  功能特性:                                                     ║
    ║  • 从 my_cookie.txt 文件读取Cookie                            ║
    ║  • 支持选择单本或多本作品备份                                  ║
    ║  • 实时保存，边下载边生成DOCX文件                             ║
    ║  • 完整保留正文格式和作者有话说                               ║
    ║  • 自动添加章节编号和层级标题                                 ║
    ║                                                                ║
    ║  使用方法:                                                     ║
    ║  1. 准备Cookie: 登录晋江→F12→Network→复制Cookie到txt文件     ║
    ║  2. 运行程序: python jjwxc_col.py                            ║
    ║  3. 选择作品: 根据提示选择要备份的作品                        ║
    ║  4. 查看结果: 备份完成后查看生成的DOCX文件                   ║
    ╚════════════════════════════════════════════════════════════════╝
    """)
    
    # 检查Cookie文件
    if not os.path.exists(COOKIE_FILE):
        print(f"❌ 未找到 {COOKIE_FILE} 文件")
        print("\n📝 Cookie获取步骤:")
        print("1. 使用浏览器登录晋江文学城作者后台")
        print("2. 按F12打开开发者工具")
        print("3. 切换到Network(网络)选项卡")
        print("4. 刷新页面，点击任意请求")
        print("5. 在Request Headers中找到'Cookie'字段")
        print("6. 复制完整的Cookie值")
        print(f"7. 创建 {COOKIE_FILE} 文件，粘贴Cookie内容并保存")
        print("\n按回车键退出...")
        input()
        exit(1)
    
    # 启动备份工具
    try:
        tool = JJWXCBackupTool()
        tool.backup_all_novels()
    except KeyboardInterrupt:
        print("\n\n用户中断程序")
    except Exception as e:
        print(f"\n程序运行出错: {e}")
        print("请检查网络连接和Cookie是否有效")
    finally:
        print("\n程序结束，按回车键退出...")
        input()