#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
=================================================================
                     DOCX格式输出测试
=================================================================
功能：测试DOCX文档生成和格式保留功能

使用场景：
- 验证DOCX文档创建是否正常
- 检查换行符和空行保留
- 调试文档格式和样式

测试内容：
- 使用测试内容创建DOCX文档
- 验证段落结构和换行保留
- 检查作者有话说格式化
- 确认空行处理效果

注意：生成的测试文档会保存到tests目录
=================================================================
"""
import os
import sys
# 添加父目录到Python路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from jjwxc_col import JJWXCBackupTool
from docx import Document
from docx.shared import RGBColor

def test_docx_format():
    """测试DOCX格式输出"""
    
    # 初始化工具
    tool = JJWXCBackupTool()
    
    print("=" * 60)
    print("DOCX格式输出测试")
    print("=" * 60)
    
    # 创建测试内容
    test_content = """第一行正文内容
第二行正文内容

第四行正文内容（前面有空行）
第五行正文内容
第六行正文内容

最后一行正文内容【作者有话说】
作者备注第一行
作者备注第二行

作者备注第四行（前面有空行）
作者备注最后一行"""
    
    print("测试内容准备:")
    print(f"内容长度: {len(test_content)} 字符")
    print(f"内容行数: {len(test_content.split(chr(10)))}")
    print(f"包含作者有话说: {'是' if '【作者有话说】' in test_content else '否'}")
    
    # 创建DOCX文档
    print("\n创建DOCX文档...")
    doc = Document()
    
    # 添加测试标题
    title = doc.add_heading('DOCX格式测试文档', level=0)
    title.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    # 添加章节标题
    chapter_title = doc.add_heading('第一章 测试章节', level=1)
    
    # 使用工具的内容添加方法
    tool._add_content_to_doc(doc, test_content)
    
    # 保存文档
    output_path = os.path.join("tests", "docx_format_test.docx")
    doc.save(output_path)
    
    print(f"✓ DOCX文档已保存: {output_path}")
    
    # 读取并验证文档内容
    print("\n验证生成的文档...")
    saved_doc = Document(output_path)
    
    print(f"文档段落总数: {len(saved_doc.paragraphs)}")
    
    # 分析段落内容
    title_paragraphs = []
    content_paragraphs = []
    empty_paragraphs = []
    
    for i, paragraph in enumerate(saved_doc.paragraphs):
        text = paragraph.text
        if paragraph.style.name.startswith('Heading'):
            title_paragraphs.append((i, text))
        elif text == "":
            empty_paragraphs.append(i)
        else:
            content_paragraphs.append((i, text))
    
    print(f"标题段落: {len(title_paragraphs)}个")
    for i, text in title_paragraphs:
        print(f"  段落{i}: {text}")
    
    print(f"内容段落: {len(content_paragraphs)}个")
    print(f"空段落(空行): {len(empty_paragraphs)}个，位置: {empty_paragraphs}")
    
    # 检查是否有作者有话说标题
    author_note_found = any('作者有话说' in text for _, text in title_paragraphs)
    print(f"作者有话说标题: {'找到' if author_note_found else '未找到'}")
    
    # 显示前10个内容段落作为预览
    print("\n内容预览（前10个段落）:")
    print("-" * 40)
    for i, (para_idx, text) in enumerate(content_paragraphs[:10]):
        print(f"段落{para_idx}: '{text}'")
    
    if len(content_paragraphs) > 10:
        print(f"... (还有{len(content_paragraphs)-10}个内容段落)")
    
    print("\n✓ DOCX格式测试完成")
    print(f"可以打开 {output_path} 查看实际效果")

if __name__ == "__main__":
    test_docx_format()
